"""Generate the DunongAI Activity 3 Questionnaire as a Word (.docx) document.

Usage:
    pip install python-docx
    python generate_questionnaire.py

Output:
    DunongAI_Activity3_Questionnaire.docx (in the same directory)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


LIKERT_LABELS = ["1 - Strongly Disagree", "2 - Disagree", "3 - Neutral",
                 "4 - Agree", "5 - Strongly Agree"]


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p


def add_likert_question(doc, number, question_text):
    """Adds a numbered Likert-scale question with a 1-5 grid below it."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{number}. {question_text}")
    run.font.size = Pt(11)

    table = doc.add_table(rows=2, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, label in enumerate(LIKERT_LABELS):
        header = table.cell(0, i)
        header.text = label
        for paragraph in header.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in paragraph.runs:
                r.font.size = Pt(9)
                r.font.bold = True
        # Empty checkbox row
        check = table.cell(1, i)
        check.text = "☐"
        for paragraph in check.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in paragraph.runs:
                r.font.size = Pt(14)
    doc.add_paragraph()


def add_open_question(doc, number, question_text, lines=3):
    """Adds a numbered open-ended question with blank lines for the answer."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{number}. {question_text}")
    run.font.size = Pt(11)

    for _ in range(lines):
        line = doc.add_paragraph("_" * 95)
        line.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def add_demographic_field(doc, label, options=None):
    p = doc.add_paragraph()
    run = p.add_run(f"   • {label}: ")
    run.bold = True
    if options:
        p.add_run("  ".join([f"☐ {opt}" for opt in options]))
    else:
        p.add_run("_" * 60)


def add_section_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_hat_header(doc, hat_name, hat_color_hex, description):
    """Adds a styled section header for each of the 6 Thinking Hats."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, hat_color_hex)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(hat_name)
    run.bold = True
    run.font.size = Pt(16)
    # White text for darker hats, black for lighter
    if hat_color_hex in ("000000", "0070C0", "00B050"):
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    else:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(description)
    r2.italic = True
    r2.font.size = Pt(10)
    doc.add_paragraph()


def build_document():
    doc = Document()

    # ============ DOCUMENT STYLES ============
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============ COVER PAGE ============
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DunongAI")
    run.font.size = Pt(48)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("AI-Powered Adaptive Reading Companion for Filipino Learners")
    sub_run.italic = True
    sub_run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    activity = doc.add_paragraph()
    activity.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a_run = activity.add_run("Activity 3 — Survey Questionnaire")
    a_run.bold = True
    a_run.font.size = Pt(20)

    subj = doc.add_paragraph()
    subj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subj.add_run("Design Thinking & 6 Thinking Hats Assessment")
    s_run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fields = [
        ("Respondent Name:", ""),
        ("Date:", ""),
        ("School / Organization:", ""),
        ("Role:", "☐ Student   ☐ Parent   ☐ Teacher   ☐ Admin"),
    ]
    for i, (label, value) in enumerate(fields):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value if value else "_" * 40
        for cell in info_table.row_cells(i):
            for paragraph in cell.paragraphs:
                for r in paragraph.runs:
                    r.font.size = Pt(11)
                    if cell == info_table.cell(i, 0):
                        r.bold = True

    add_section_break(doc)

    # ============ INTRODUCTION ============
    add_heading(doc, "Introduction & Informed Consent", level=1)

    add_paragraph(doc,
        "Thank you for taking the time to participate in this survey. "
        "DunongAI is an AI-powered adaptive reading companion designed for K-6 "
        "Filipino public school students. It features adaptive level testing, "
        "a Basa Bot vocabulary helper, AI-generated comprehension questions, "
        "and a teacher dashboard with class analytics — built to support reading "
        "proficiency in low-connectivity environments.")

    add_paragraph(doc,
        "This questionnaire is part of an academic technopreneurship study "
        "(Activity 3 — Design Thinking and 6 Thinking Hats). Your honest feedback "
        "will help us evaluate the prototype, identify strengths and weaknesses, "
        "and improve the product before public release.")

    add_paragraph(doc, "Survey Details:", bold=True)
    bullets = [
        "Estimated completion time: 10–15 minutes",
        "Number of sections: 5 (Demographics, Empathy, 6 Thinking Hats, Synthesis, SWOT)",
        "Your responses are CONFIDENTIAL and will be used only for academic purposes.",
        "Participation is voluntary; you may skip any question or withdraw at any time.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    add_paragraph(doc, "Likert Scale Used Throughout this Survey:", bold=True)
    doc.add_paragraph(
        "1 = Strongly Disagree    2 = Disagree    3 = Neutral    "
        "4 = Agree    5 = Strongly Agree",
        style='Intense Quote')

    add_paragraph(doc,
        "Please answer all questions to the best of your ability. "
        "For Section C (6 Thinking Hats), you will be shown a brief walkthrough "
        "or screenshots of the DunongAI prototype before answering.", italic=True)

    add_section_break(doc)

    # ============ SECTION A: DEMOGRAPHICS ============
    add_heading(doc, "Section A — Demographic Profile", level=1)
    add_paragraph(doc,
        "Please answer the following questions about yourself. "
        "All information is kept confidential.", italic=True)

    add_demographic_field(doc, "Respondent Type",
                          ["Student (K-6)", "Parent / Guardian", "Teacher", "School Admin"])
    add_demographic_field(doc, "Age Range",
                          ["Under 12", "12-17", "18-25", "26-35", "36-50", "Over 50"])
    add_demographic_field(doc, "Sex", ["Male", "Female", "Prefer not to say"])
    add_demographic_field(doc, "Grade Level (if Student)",
                          ["K", "1", "2", "3", "4", "5", "6"])
    add_demographic_field(doc, "Years of Teaching (if Teacher)",
                          ["0-2", "3-5", "6-10", "11-20", "20+"])
    add_demographic_field(doc, "Type of School",
                          ["Public Elementary", "Private Elementary", "Other"])
    add_demographic_field(doc, "Region / Province")
    add_demographic_field(doc, "Devices Available at Home/School",
                          ["Smartphone", "Tablet", "Laptop/PC", "None"])
    add_demographic_field(doc, "Internet Reliability",
                          ["Always available", "Sometimes", "Rarely", "Never"])
    add_demographic_field(doc, "Familiarity with educational apps",
                          ["Very familiar", "Somewhat", "Not at all"])

    add_section_break(doc)

    # ============ SECTION B: EMPATHY QUESTIONS ============
    add_heading(doc, "Section B — Empathy Questions (Current Situation Without DunongAI)", level=1)
    add_paragraph(doc,
        "These questions explore your current experiences and pain points "
        "with reading instruction, support, or learning BEFORE using DunongAI.",
        italic=True)

    # --- B1: For Students / Parents ---
    add_heading(doc, "B1. For Students / Parents (Answer on behalf of the child)", level=2)
    b1_questions = [
        "My child / I struggle to understand stories written in Filipino.",
        "My child / I often encounter unfamiliar Filipino words while reading.",
        "Traditional textbooks feel boring or unengaging.",
        "It is difficult to find reading materials matched to my child's level.",
        "Helping my child with reading takes more time than I can give.",
        "I (or my child) feel discouraged when reading is hard.",
        "I do not have an easy way to track my child's reading progress.",
        "I (or my child) avoid reading because it feels stressful.",
    ]
    for i, q in enumerate(b1_questions, start=1):
        add_likert_question(doc, i, q)

    add_open_question(doc, len(b1_questions) + 1,
                      "What is the biggest difficulty you or your child faces with reading today?", lines=3)
    add_open_question(doc, len(b1_questions) + 2,
                      "What tools or resources (books, apps, tutors) do you currently use for reading practice?", lines=3)

    # --- B2: For Teachers ---
    add_heading(doc, "B2. For Teachers", level=2)
    b2_questions = [
        "Assessing each student's individual reading level is time-consuming.",
        "I find it difficult to provide differentiated instruction in a large class.",
        "Identifying struggling readers early in the school year is challenging.",
        "Tracking each student's reading progress involves significant manual work.",
        "My current tools do not give me clear data on class-wide reading proficiency.",
        "I have limited time per student for one-on-one reading support.",
        "Communicating reading progress to parents is difficult or infrequent.",
    ]
    for i, q in enumerate(b2_questions, start=1):
        add_likert_question(doc, i, q)

    add_open_question(doc, len(b2_questions) + 1,
                      "What is your biggest pain point in teaching reading today?", lines=3)
    add_open_question(doc, len(b2_questions) + 2,
                      "What kind of data or tool would make your job easier?", lines=3)

    # --- B3: For School Admins ---
    add_heading(doc, "B3. For School Administrators", level=2)
    b3_questions = [
        "Our school struggles with reading proficiency at the K-6 level.",
        "We lack data-driven insights into reading performance across classes.",
        "Budget constraints limit investment in new ed-tech tools.",
        "Teacher training time is a barrier to adopting new tools.",
        "Parents often request more visibility into their child's academic progress.",
    ]
    for i, q in enumerate(b3_questions, start=1):
        add_likert_question(doc, i, q)

    add_open_question(doc, len(b3_questions) + 1,
                      "What is your school's biggest challenge regarding K-6 reading outcomes?", lines=3)

    add_section_break(doc)

    # ============ SECTION C: 6 THINKING HATS ============
    add_heading(doc, "Section C — 6 Thinking Hats Assessment of DunongAI", level=1)
    add_paragraph(doc,
        "Before answering this section, you should have been shown a short demo "
        "or screenshots of the DunongAI prototype (Story Library, Basa Bot, "
        "Comprehension Questions, Teacher Dashboard, and Progress Tracking). "
        "Each 'hat' below represents a different lens through which you will "
        "evaluate the app.", italic=True)

    # ---- WHITE HAT ----
    add_hat_header(doc, "⚪ WHITE HAT — Facts & Information",
                   "FFFFFF",
                   "Focus on objective facts, data, accuracy, and information the app provides.")
    white_questions = [
        ("likert", "The diagnostic quiz appears to accurately assess a child's reading level."),
        ("likert", "The reading levels (Grades 1-6) match the difficulty I would expect for each grade."),
        ("likert", "The Filipino story library is varied enough for daily use over a school year."),
        ("likert", "The AI-generated comprehension questions are clearly written and aligned with the story."),
        ("likert", "The Basa Bot vocabulary helper provides accurate word definitions."),
        ("likert", "The teacher dashboard presents the right information for instructional decisions."),
        ("open",   "Approximately how many minutes per day would a child realistically use this app?"),
        ("open",   "What additional information or data would you want the app to display?"),
    ]
    for i, (qtype, qtext) in enumerate(white_questions, start=1):
        if qtype == "likert":
            add_likert_question(doc, i, qtext)
        else:
            add_open_question(doc, i, qtext, lines=2)

    # ---- RED HAT ----
    add_hat_header(doc, "🔴 RED HAT — Feelings & Emotions",
                   "C00000",
                   "Focus on feelings, intuition, emotional reactions — no need to justify them.")
    red_questions = [
        ("likert", "I feel excited about the idea of using DunongAI in my classroom or home."),
        ("likert", "Basa Bot feels friendly and approachable to a child."),
        ("likert", "Earning badges and streaks would motivate me (or my child) to read more."),
        ("likert", "The visual design feels appropriate and appealing for Filipino children."),
        ("likert", "The app respects student dignity (does not shame children for low scores)."),
        ("likert", "I feel hopeful that this app could help children who struggle with reading."),
        ("likert", "The leaderboard creates a fun and positive sense of friendly competition."),
        ("open",   "Describe in 1-2 sentences how you felt while exploring the DunongAI prototype."),
    ]
    for i, (qtype, qtext) in enumerate(red_questions, start=1):
        if qtype == "likert":
            add_likert_question(doc, i, qtext)
        else:
            add_open_question(doc, i, qtext, lines=2)

    # ---- BLACK HAT ----
    add_hat_header(doc, "⚫ BLACK HAT — Caution & Critical Judgment",
                   "000000",
                   "Identify risks, weaknesses, potential problems, and reasons for caution.")
    black_questions = [
        ("likert", "I am concerned about the screen-time impact on young children."),
        ("likert", "I worry that the AI may sometimes give inaccurate definitions or feedback."),
        ("likert", "Data privacy of student information is a significant concern."),
        ("likert", "Leaderboards may put unhealthy pressure on struggling readers."),
        ("likert", "Schools with weak or no internet will not be able to fully benefit."),
        ("likert", "Over-reliance on AI may reduce teachers' own pedagogical judgment."),
        ("likert", "Older children may find the app's design 'too childish' or vice-versa."),
        ("open",   "What is the BIGGEST risk or downside you see with DunongAI?"),
    ]
    for i, (qtype, qtext) in enumerate(black_questions, start=1):
        if qtype == "likert":
            add_likert_question(doc, i, qtext)
        else:
            add_open_question(doc, i, qtext, lines=3)

    # ---- YELLOW HAT ----
    add_hat_header(doc, "🟡 YELLOW HAT — Benefits & Optimism",
                   "FFC000",
                   "Focus on benefits, value, optimism, and reasons this idea will work.")
    yellow_questions = [
        ("likert", "DunongAI would save teachers significant time on student assessment."),
        ("likert", "Personalized reading levels will improve student learning outcomes."),
        ("likert", "Real-time word definitions (Basa Bot) is a genuinely valuable feature."),
        ("likert", "Filipino-language content makes this app stand out vs. international apps."),
        ("likert", "AI insights would help teachers identify struggling readers earlier."),
        ("likert", "Demo mode (offline) makes the app usable in low-connectivity schools."),
        ("likert", "Parents would appreciate visibility into their child's reading progress."),
        ("open",   "Which benefit would have the BIGGEST positive impact on Filipino public schools?"),
    ]
    for i, (qtype, qtext) in enumerate(yellow_questions, start=1):
        if qtype == "likert":
            add_likert_question(doc, i, qtext)
        else:
            add_open_question(doc, i, qtext, lines=3)

    # ---- GREEN HAT ----
    add_hat_header(doc, "🟢 GREEN HAT — Creativity & New Ideas",
                   "00B050",
                   "Generate new ideas, alternatives, and creative improvements.")
    green_questions = [
        ("likert", "Adding audio narration / text-to-speech would improve the app."),
        ("likert", "Supporting regional dialects (Cebuano, Ilocano, Hiligaynon) is important."),
        ("likert", "A parent-facing module showing weekly progress would add real value."),
        ("likert", "Reading clubs or multiplayer story modes would engage children more."),
        ("likert", "Customizable avatars or themed environments would increase engagement."),
        ("open",   "Suggest ONE new feature you would add to DunongAI."),
        ("open",   "How could the reward system (badges/streaks/leaderboards) be improved?"),
        ("open",   "What partnerships (DepEd, NGOs, publishers) should DunongAI pursue?"),
    ]
    for i, (qtype, qtext) in enumerate(green_questions, start=1):
        if qtype == "likert":
            add_likert_question(doc, i, qtext)
        else:
            add_open_question(doc, i, qtext, lines=3)

    # ---- BLUE HAT ----
    add_hat_header(doc, "🔵 BLUE HAT — Process & Overview",
                   "0070C0",
                   "Big-picture thinking: rollout, pricing, success metrics, and overall strategy.")
    blue_questions = [
        ("likert", "DunongAI fits well within the current DepEd K-6 curriculum."),
        ("likert", "Teachers would need only minimal training to adopt this app."),
        ("likert", "A freemium model (free for students, paid for schools) seems fair."),
        ("likert", "Rolling out per grade level is better than school-wide all at once."),
        ("likert", "Success should be measured by reading proficiency gains, not just usage."),
        ("likert", "Overall, I would recommend DunongAI to my school or community."),
        ("open",   "What is a FAIR monthly/yearly price per student or per teacher (in PHP)?"),
        ("open",   "What is your OVERALL recommendation for DunongAI's next steps?"),
    ]
    for i, (qtype, qtext) in enumerate(blue_questions, start=1):
        if qtype == "likert":
            add_likert_question(doc, i, qtext)
        else:
            add_open_question(doc, i, qtext, lines=3)

    add_section_break(doc)

    # ============ SECTION D: OPEN-ENDED SYNTHESIS ============
    add_heading(doc, "Section D — Final Open-Ended Synthesis", level=1)
    add_paragraph(doc,
        "Three final open-ended questions to capture your overall reflections.",
        italic=True)

    add_open_question(doc, 1,
                      "Would you recommend DunongAI to colleagues, parents, or your school? Why or why not?",
                      lines=4)
    add_open_question(doc, 2,
                      "What is the SINGLE change that would make you most likely to adopt DunongAI?",
                      lines=4)
    add_open_question(doc, 3,
                      "Any final comments, suggestions, or concerns you wish to share?",
                      lines=5)

    add_section_break(doc)

    # ============ SECTION E: SWOT ANALYSIS ============
    add_heading(doc, "Section E — SWOT Analysis (Researcher's Predicted Assessment)", level=1)
    add_paragraph(doc,
        "Below is a predicted SWOT analysis of DunongAI based on the prototype's "
        "current feature set. This serves as a baseline; survey results will be "
        "used to refine and validate this analysis in Week 10.", italic=True)

    swot_table = doc.add_table(rows=3, cols=2)
    swot_table.style = 'Light Grid Accent 1'
    swot_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    h1 = swot_table.cell(0, 0)
    h1.text = "STRENGTHS (Internal — Positive)"
    set_cell_shading(h1, "C6E0B4")
    h2 = swot_table.cell(0, 1)
    h2.text = "WEAKNESSES (Internal — Negative)"
    set_cell_shading(h2, "F8CBAD")
    for cell in [h1, h2]:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(12)

    strengths = [
        "Filipino-first content — unique vs. Khan Academy, Duolingo, etc.",
        "Adaptive AI leveling — no manual setup required from teacher",
        "Real-time Basa Bot vocabulary support during reading",
        "Teacher dashboard with AI-generated insights and trend charts",
        "Demo mode works offline / with no API keys — ideal for low-connectivity schools",
        "Built with modern, scalable stack (React + Vite + Claude AI)",
        "Engaging gamification: badges, streaks, leaderboards",
        "Free to start, easy onboarding via diagnostic quiz",
    ]
    weaknesses = [
        "Limited story library at initial launch (curated set)",
        "Dependency on Claude API — operating costs scale with usage",
        "Narrow scope: K-6 only, no path to upper grades yet",
        "Requires teacher training/buy-in for full dashboard value",
        "No native mobile app yet (web-based React only)",
        "No multi-dialect support (Cebuano, Ilocano, Hiligaynon) at launch",
        "Limited offline functionality beyond demo mode",
        "Brand recognition is zero — competing against established names",
    ]

    s_cell = swot_table.cell(1, 0)
    s_cell.text = ""
    for item in strengths:
        p = s_cell.add_paragraph(f"• {item}")
        for r in p.runs:
            r.font.size = Pt(10)

    w_cell = swot_table.cell(1, 1)
    w_cell.text = ""
    for item in weaknesses:
        p = w_cell.add_paragraph(f"• {item}")
        for r in p.runs:
            r.font.size = Pt(10)

    # Second header row
    h3 = swot_table.cell(2, 0)
    h3.text = ""
    set_cell_shading(h3, "BDD7EE")
    h3_p = h3.paragraphs[0]
    h3_r = h3_p.add_run("OPPORTUNITIES (External — Positive)")
    h3_r.bold = True
    h3_r.font.size = Pt(12)
    h4 = swot_table.cell(2, 1)
    h4.text = ""
    set_cell_shading(h4, "FFE699")
    h4_p = h4.paragraphs[0]
    h4_r = h4_p.add_run("THREATS (External — Negative)")
    h4_r.bold = True
    h4_r.font.size = Pt(12)

    # Add an extra row for opportunities/threats content
    swot_table.add_row()
    opps_cell = swot_table.cell(3, 0)
    threats_cell = swot_table.cell(3, 1)

    opportunities = [
        "Partnership with DepEd for public school pilot programs",
        "Expansion into regional dialects (Cebuano, Ilocano, Hiligaynon)",
        "Parent-facing freemium tier for home use",
        "Funding from NGOs and literacy-focused foundations",
        "Integration with Google Classroom / existing LMS platforms",
        "Teacher professional development (PD) packages as added revenue",
        "Partnerships with Filipino publishers for licensed story content",
        "Government push for digital learning post-pandemic",
    ]
    threats = [
        "Data Privacy Act (DPA 2012) compliance burden",
        "Competition from international ed-tech (Khan Kids, Duolingo Kids)",
        "Limited ed-tech budgets in public schools",
        "Unreliable internet infrastructure in rural areas",
        "AI hallucinations may damage trust if not carefully managed",
        "Screen-time backlash from parents and policymakers",
        "Claude/Anthropic API pricing changes affecting unit economics",
        "Possible DepEd procurement / accreditation barriers",
    ]

    opps_cell.text = ""
    for item in opportunities:
        p = opps_cell.add_paragraph(f"• {item}")
        for r in p.runs:
            r.font.size = Pt(10)

    threats_cell.text = ""
    for item in threats:
        p = threats_cell.add_paragraph(f"• {item}")
        for r in p.runs:
            r.font.size = Pt(10)

    doc.add_paragraph()
    add_paragraph(doc, "Researcher's Strategic Insight:", bold=True)
    add_paragraph(doc,
        "DunongAI's strongest competitive moat is Filipino-language, "
        "culturally-relevant content combined with adaptive AI assessment — "
        "a combination international competitors cannot easily replicate. "
        "The greatest near-term risk is API cost economics and adoption "
        "barriers in low-budget public schools. Survey responses will help "
        "validate whether the predicted weaknesses and threats align with "
        "actual user perceptions.", italic=True)

    add_section_break(doc)

    # ============ THANK YOU PAGE ============
    doc.add_paragraph()
    doc.add_paragraph()
    thanks = doc.add_paragraph()
    thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = thanks.add_run("Maraming Salamat Po!")
    t_run.bold = True
    t_run.font.size = Pt(32)
    t_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub_thanks = doc.add_paragraph()
    sub_thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st_run = sub_thanks.add_run("Thank you for your valuable contribution to this study.")
    st_run.italic = True
    st_run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(
        "Your responses will be analyzed and the results will inform the next "
        "iteration of DunongAI. Survey results will be presented in Week 10."
    ).font.size = Pt(11)

    return doc


if __name__ == "__main__":
    doc = build_document()
    output_path = "DunongAI_Activity3_Questionnaire.docx"
    doc.save(output_path)
    print(f"[OK] Generated: {output_path}")
